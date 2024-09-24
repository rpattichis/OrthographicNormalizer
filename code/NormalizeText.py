import docx
import pandas as pd

"""
This function will take the path and name of a docx and write it to a txt.
"""

class NormalizeText:
  """
  Inputs:
    - path: should be that path to find the rules
    - rules_files: a list of excel documents that contains an ordered
                    list of the files that contain the rules.
  """
  def __init__(self, config) -> None:
      self.DATA_PATH = config["data_path"]
      self.RULE_PATH = config["rules_path"]
      # converts it to a list of Pandas
      self.rules = list(map(lambda x: pd.read_excel(self.RULE_PATH + x).fillna(''), config["rules_list"]))
      
      # variables below are for removing the second stress
      self.exceptions = config["exceptions"]
      self.accent_dict = config["vowel_mappings"]

      self.BOL_token = config["BOL_token"]
      self.EOL_token = config["EOL_token"]

  def __check_length__(self, infile, outfile) -> bool:
      return len(infile) == len(outfile)
  
  def __read_word_file__(self, filename: str) -> list[str]:
      doc = docx.Document(self.DATA_PATH + filename)
      lines = []
      for para in doc.paragraphs:
          lines.append(para.text)
      return lines
  
  def __read_txt_file__(self, filename: str) -> list[str]:
      f = open(self.DATA_PATH + filename, "r")
      lines = []
      for l in f:
        lines.append(l)
      return lines

  def __doc_to_txt__(self, filename: str, lines: list[str]) -> None:
      fullTextPar = '\n'.join(lines)
      
      f = open(self.DATA_PATH + filename + '.txt', "w", encoding='utf-8')
      f.write(fullTextPar)
      f.close()

  def __lines_to_doc__(self, filename: str, lines: list[str]) -> None:
      doc = docx.Document()
      for line in lines:
        p = doc.add_paragraph()
        p.add_run(line)
        p = doc.add_paragraph()
      doc.save(self.DATA_PATH + filename + '.docx')
  
  def __stress_helper__(self, word: str) -> str:
    accents = list(filter(lambda letter: letter in self.accent_dict.keys(), word))
    if len(accents) > 1 and word not in self.exceptions:
      # use the accents list to identify the second character, and map it to its unaccented version
      # i think this has to be done by reversing the word string first (in the case that the accented
      # words found are the same character)
      # NOTE: the implementation below will definitely be expensive, look into other ways to do this
      reversed_word = word[::-1]
      accent_index = reversed_word.find(accents[1])
      reversed_word[accent_index] = self.accent_dict[accents[1]] # replace the key w the value
      return reversed_word[::-1]
    return word

  def __remove_second_stress__(self, line: str) -> str:
    # Step 1: break the line into list of words
    words = line.split()

    # Step 2: fix the words (ideally in parallel) for the second stress correction
    # i.e., this should eventually become map(self.__stress_helper__, words)
    words = list(map(self.__stress_helper__, words))
    return ' '.join(words)

  """
  The rules have to be applied in sequential order for each line.
  NOTE: Including the second stress removal, which has to happen
        after the smoothing and before the corrections rules. I will
        implement it as a helper function which is called here.
  """
  def __FaR_helper__(self, line: str) -> str:
      # apply this for each rule file
      for j, rule_list in enumerate(self.rules):
        # NOTE: this assumes we are the beginning of everything.
        if j == 0 and self.BOL_token != None:
          # we need to add specific things at the begging and end of each line
          line = self.BOL_token + line + self.EOL_token
        # NOTE: this assumes that we are right before the 'corrections' and after 
        # the 'smoothing' rules.
        elif j == 1 and self.exceptions != None:
          # we need to remove the second stress right after smoothing document
          line = self.__remove_second_stress__(line)

        find = rule_list['Find']
        replace = rule_list['Replace']

        for i in range(len(find)):
          line = line.replace(find[i], replace[i])
        
      return line
  
  def find_and_replace(self, filename: str, is_word_doc=True) -> None:
      # Step 1: Convert word to text doc is required
      lines = None
      if is_word_doc:
        lines = self.__read_word_file__(filename)
        self.__doc_to_txt__(filename, lines)
      else:
        lines = self.__read_txt_file__(filename)

      # Step 2: Apply the rules in appropriate order for all lines
      processed_lines = list(map(self.__FaR_helper__, lines))

      # Step 3: Save the updated lines into a txt and word doc version 
      filename = filename.split(".")[0] # parse the filename
      self.__doc_to_txt__(filename + '-FaR', processed_lines)
      self.__lines_to_doc__(filename + '-FaR', processed_lines)